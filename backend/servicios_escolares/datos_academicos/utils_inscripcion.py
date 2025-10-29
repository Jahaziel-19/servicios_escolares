from docxtpl import DocxTemplate
from datetime import datetime, date
from io import BytesIO
from django.http import HttpResponse
from django.conf import settings
import os
from .models_inscripcion import Inscripcion
from .models import Alumno, Carrera, PeriodoEscolar
from openpyxl import load_workbook


def generar_formato_inscripcion(inscripcion_id, plantilla_path=None):
    """
    Genera el formato de inscripción en formato DOCX usando una plantilla
    """
    try:
        inscripcion = Inscripcion.objects.get(id=inscripcion_id)
        
        # Usar plantilla por defecto si no se especifica una
        if not plantilla_path:
            plantilla_path = os.path.join(
                settings.MEDIA_ROOT, 
                'plantillas', 
                'formato_inscripcion.docx'
            )
        
        # Verificar que existe la plantilla
        if not os.path.exists(plantilla_path):
            raise FileNotFoundError(f"No se encontró la plantilla en: {plantilla_path}")
        
        # Crear contexto con datos de la inscripción
        contexto = crear_contexto_inscripcion(inscripcion)
        
        # Generar documento
        doc = DocxTemplate(plantilla_path)
        doc.render(contexto)
        
        # Crear buffer para la respuesta
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Crear nombre del archivo
        nombre_archivo = f"inscripcion_{inscripcion.folio}_{inscripcion.nombre.replace(' ', '_')}.docx"
        
        # Crear respuesta HTTP
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
        
        return response
        
    except Inscripcion.DoesNotExist:
        raise ValueError(f"No se encontró la inscripción con ID: {inscripcion_id}")
    except Exception as e:
        raise Exception(f"Error al generar formato de inscripción: {str(e)}")


def generar_formato_reinscripcion(reinscripcion_id, plantilla_path=None):
    raise NotImplementedError('Función de reinscripción eliminada')


def crear_contexto_inscripcion(inscripcion):
    """
    Crea el contexto de variables para el formato de inscripción
    """
    contexto = {
        # Información básica
        'folio': inscripcion.folio,
        'fecha_inscripcion': inscripcion.fecha_inscripcion.strftime('%d/%m/%Y'),
        'fecha_emision': datetime.now().strftime('%d/%m/%Y'),
        
        # Datos personales
        'nombre_completo': f"{inscripcion.nombre} {inscripcion.apellido_paterno} {inscripcion.apellido_materno}".strip(),
        'nombre': inscripcion.nombre,
        'apellido_paterno': inscripcion.apellido_paterno,
        'apellido_materno': inscripcion.apellido_materno,
        'curp': inscripcion.curp,
        'fecha_nacimiento': inscripcion.fecha_nacimiento.strftime('%d/%m/%Y') if inscripcion.fecha_nacimiento else '',
        'lugar_nacimiento': inscripcion.lugar_nacimiento or '',
        'genero': inscripcion.get_genero_display(),
        'estado_civil': inscripcion.get_estado_civil_display(),
        
        # Datos académicos
        'carrera': inscripcion.carrera.nombre if inscripcion.carrera else '',
        'carrera_clave': inscripcion.carrera.clave if inscripcion.carrera else '',
        'periodo_escolar': str(inscripcion.periodo_escolar) if inscripcion.periodo_escolar else '',
        'semestre': inscripcion.semestre,
        'modalidad': inscripcion.get_modalidad_display(),
        'turno': inscripcion.get_turno_display(),
        
        # Datos de contacto
        'telefono': inscripcion.telefono or '',
        'email': inscripcion.email or '',
        'telefono_emergencia': inscripcion.telefono_emergencia or '',
        'contacto_emergencia': inscripcion.contacto_emergencia or '',
        
        # Dirección
        'calle': inscripcion.calle or '',
        'numero': inscripcion.numero or '',
        'colonia': inscripcion.colonia or '',
        'municipio': inscripcion.municipio or '',
        'estado': inscripcion.estado or '',
        'codigo_postal': inscripcion.codigo_postal or '',
        
        # Datos académicos previos
        'escuela_procedencia': inscripcion.escuela_procedencia or '',
        'promedio_bachillerato': str(inscripcion.promedio_bachillerato) if inscripcion.promedio_bachillerato else '',
        'año_egreso': str(inscripcion.año_egreso) if inscripcion.año_egreso else '',
        
        # Documentos
        'acta_nacimiento': 'Sí' if inscripcion.acta_nacimiento else 'No',
        'certificado_bachillerato': 'Sí' if inscripcion.certificado_bachillerato else 'No',
        'curp_documento': 'Sí' if inscripcion.curp_documento else 'No',
        'fotografias': 'Sí' if inscripcion.fotografias else 'No',
        'certificado_medico': 'Sí' if inscripcion.certificado_medico else 'No',
        
        # Estado del proceso
        'estatus': inscripcion.get_estatus_display(),
        'observaciones': inscripcion.observaciones or '',
    }
    
    return contexto


def crear_contexto_reinscripcion(reinscripcion):
    return {}


def crear_plantillas_por_defecto():
    """
    Crea las plantillas por defecto para inscripción si no existen
    """
    plantillas_dir = os.path.join(settings.MEDIA_ROOT, 'plantillas')
    os.makedirs(plantillas_dir, exist_ok=True)
    plantilla_inscripcion = os.path.join(plantillas_dir, 'formato_inscripcion.docx')
    if not os.path.exists(plantilla_inscripcion):
        crear_plantilla_inscripcion_basica(plantilla_inscripcion)


def crear_plantilla_inscripcion_basica(ruta_archivo):
    """
    Crea una plantilla básica de inscripción con variables de ejemplo
    """
    from docx import Document
    from docx.shared import Inches
    
    doc = Document()
    
    # Título
    titulo = doc.add_heading('FORMATO DE INSCRIPCIÓN', 0)
    titulo.alignment = 1  # Centrado
    
    # Información básica
    doc.add_heading('INFORMACIÓN BÁSICA', level=1)
    doc.add_paragraph(f'Folio: {{{{ folio }}}}')
    doc.add_paragraph(f'Fecha de inscripción: {{{{ fecha_inscripcion }}}}')
    doc.add_paragraph(f'Período escolar: {{{{ periodo_escolar }}}}')
    
    # Datos personales
    doc.add_heading('DATOS PERSONALES', level=1)
    doc.add_paragraph(f'Nombre completo: {{{{ nombre_completo }}}}')
    doc.add_paragraph(f'CURP: {{{{ curp }}}}')
    doc.add_paragraph(f'Fecha de nacimiento: {{{{ fecha_nacimiento }}}}')
    doc.add_paragraph(f'Lugar de nacimiento: {{{{ lugar_nacimiento }}}}')
    doc.add_paragraph(f'Género: {{{{ genero }}}}')
    doc.add_paragraph(f'Estado civil: {{{{ estado_civil }}}}')
    
    # Datos académicos
    doc.add_heading('DATOS ACADÉMICOS', level=1)
    doc.add_paragraph(f'Carrera: {{{{ carrera }}}}')
    doc.add_paragraph(f'Semestre: {{{{ semestre }}}}')
    doc.add_paragraph(f'Modalidad: {{{{ modalidad }}}}')
    doc.add_paragraph(f'Turno: {{{{ turno }}}}')
    doc.add_paragraph(f'Escuela de procedencia: {{{{ escuela_procedencia }}}}')
    doc.add_paragraph(f'Promedio de bachillerato: {{{{ promedio_bachillerato }}}}')
    
    # Datos de contacto
    doc.add_heading('DATOS DE CONTACTO', level=1)
    doc.add_paragraph(f'Teléfono: {{{{ telefono }}}}')
    doc.add_paragraph(f'Email: {{{{ email }}}}')
    doc.add_paragraph(f'Contacto de emergencia: {{{{ contacto_emergencia }}}}')
    doc.add_paragraph(f'Teléfono de emergencia: {{{{ telefono_emergencia }}}}')
    
    # Dirección
    doc.add_heading('DIRECCIÓN', level=1)
    doc.add_paragraph(f'Calle y número: {{{{ calle }}}} {{{{ numero }}}}')
    doc.add_paragraph(f'Colonia: {{{{ colonia }}}}')
    doc.add_paragraph(f'Municipio: {{{{ municipio }}}}')
    doc.add_paragraph(f'Estado: {{{{ estado }}}}')
    doc.add_paragraph(f'Código postal: {{{{ codigo_postal }}}}')
    
    # Documentos
    doc.add_heading('DOCUMENTOS ENTREGADOS', level=1)
    doc.add_paragraph(f'Acta de nacimiento: {{{{ acta_nacimiento }}}}')
    doc.add_paragraph(f'Certificado de bachillerato: {{{{ certificado_bachillerato }}}}')
    doc.add_paragraph(f'CURP: {{{{ curp_documento }}}}')
    doc.add_paragraph(f'Fotografías: {{{{ fotografias }}}}')
    doc.add_paragraph(f'Certificado médico: {{{{ certificado_medico }}}}')
    
    # Observaciones
    doc.add_heading('OBSERVACIONES', level=1)
    doc.add_paragraph(f'{{{{ observaciones }}}}')
    
    # Firmas
    doc.add_paragraph('\n\n')
    doc.add_paragraph('_' * 30 + '                    ' + '_' * 30)
    doc.add_paragraph('Firma del alumno                           Firma del responsable')
    
    doc.save(ruta_archivo)


def crear_plantilla_reinscripcion_basica(ruta_archivo):
    pass


# Módulo: utils_inscripcion (añadir funciones)
def validar_documentos_y_pagos(matricula: str, folio_pago: str = '', referencia: str = '') -> dict:
    """
    Valida documentos y pagos contra archivos Excel en la carpeta de producto no conforme.
    Busca coincidencias por matrícula y/o referencias/folios.
    """
    base_dir = r"C:\Users\Jahaziel\Documents\GitHub\servicios_escolares\backend\servicios_escolares\base de datos\Procesos\producto no conforme"
    resultados = {'valido': True, 'errores': [], 'detalles': []}

    try:
        for nombre in os.listdir(base_dir):
            if not nombre.lower().endswith(('.xlsx', '.xlsm')):
                continue
            wb = load_workbook(os.path.join(base_dir, nombre), data_only=True)
            for ws in wb.worksheets:
                # Escanear filas buscando problemas
                for row in ws.iter_rows(values_only=True):
                    if not row:
                        continue
                    row_str = ' '.join([str(c) for c in row if c is not None]).lower()
                    if matricula and matricula.lower() in row_str:
                        resultados['valido'] = False
                        resultados['errores'].append('Producto no conforme asociado a la matrícula')
                        resultados['detalles'].append({'archivo': nombre, 'hoja': ws.title, 'fila': row})
                    if folio_pago and folio_pago.lower() in row_str:
                        resultados['valido'] = False
                        resultados['errores'].append('Folio de pago observado en producto no conforme')
                        resultados['detalles'].append({'archivo': nombre, 'hoja': ws.title, 'fila': row})
                    if referencia and referencia.lower() in row_str:
                        resultados['valido'] = resultados['valido'] and False
                        resultados['errores'].append('Referencia observada en producto no conforme')
                        resultados['detalles'].append({'archivo': nombre, 'hoja': ws.title, 'fila': row})
        return resultados
    except Exception as e:
        return {'valido': False, 'errores': [f'Error al validar: {str(e)}'], 'detalles': []}