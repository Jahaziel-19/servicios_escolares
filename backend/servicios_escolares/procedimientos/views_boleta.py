from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView, DetailView, CreateView
from django.http import HttpResponse, JsonResponse
from django.utils import timezone
from django.db.models import Q
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from datos_academicos.models import Alumno, PeriodoEscolar, Calificacion, MateriaCarrera
from procedimientos.models import Boleta
from docsbuilder.utils import armar_contexto_para_alumno
from collections import defaultdict
from datetime import datetime
import tempfile
import os
# Importaciones para PDF
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from io import BytesIO


class BoletaListView(LoginRequiredMixin, ListView):
    model = Boleta
    template_name = 'procedimientos/boleta_list.html'
    context_object_name = 'boletas'
    paginate_by = 10

    def get_queryset(self):
        queryset = Boleta.objects.select_related('alumno', 'periodo_escolar').all().order_by('-fecha_generacion')

        q = self.request.GET.get('q', '').strip()
        periodo = self.request.GET.get('periodo', '').strip()

        if q:
            queryset = queryset.filter(
                Q(alumno__nombre__icontains=q) |
                Q(alumno__apellido_paterno__icontains=q) |
                Q(alumno__apellido_materno__icontains=q) |
                Q(alumno__matricula__icontains=q)
            )

        if periodo:
            queryset = queryset.filter(periodo_escolar_id=periodo)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['periodos'] = PeriodoEscolar.objects.all().order_by('-año', '-ciclo')
        
        # Agregar plantillas disponibles para boletas
        from docsbuilder.models import Plantilla
        context['plantillas_boleta'] = Plantilla.objects.all()
        
        return context


@login_required
def generar_boleta_view(request):
    """Vista para mostrar el formulario de generación de boletas"""
    alumnos = Alumno.objects.all().order_by('apellido_paterno', 'apellido_materno', 'nombre')
    periodos = PeriodoEscolar.objects.all().order_by('-año', '-ciclo')
    
    context = {
        'alumnos': alumnos,
        'periodos': periodos,
    }
    
    return render(request, 'procedimientos/generar_boleta.html', context)


@login_required
def generar_boleta_documento(request, alumno_id, periodo_id):
    """Vista para generar el documento de boleta de calificaciones en Word"""
    alumno = get_object_or_404(Alumno, id=alumno_id)
    periodo = get_object_or_404(PeriodoEscolar, id=periodo_id)
    
    # Crear o obtener la boleta
    boleta, created = Boleta.objects.get_or_create(
        alumno=alumno,
        periodo_escolar=periodo,
        defaults={'generado_por': request.user}
    )
    
    # Generar el documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Encabezado con logos y título
    crear_encabezado(doc, alumno, periodo)
    
    # Información del alumno
    crear_info_alumno(doc, alumno, periodo)
    
    # Tabla de calificaciones
    crear_tabla_calificaciones(doc, boleta)
    
    # Leyenda de niveles de desempeño
    crear_leyenda_desempeno(doc)
    
    # Pie con firma
    crear_pie_firma(doc)
    
    # Guardar documento en memoria
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Boleta_{alumno.matricula}_{periodo}.docx"'
    
    doc.save(response)
    return response


@login_required
def generar_boleta_pdf(request, alumno_id, periodo_id):
    """Vista para generar el documento de boleta de calificaciones en PDF"""
    alumno = get_object_or_404(Alumno, id=alumno_id)
    periodo = get_object_or_404(PeriodoEscolar, id=periodo_id)
    
    # Crear o obtener la boleta
    boleta, created = Boleta.objects.get_or_create(
        alumno=alumno,
        periodo_escolar=periodo,
        defaults={'generado_por': request.user}
    )
    
    # Crear buffer para el PDF
    buffer = BytesIO()
    
    # Crear el documento PDF
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Obtener estilos
    styles = getSampleStyleSheet()
    
    # Crear estilos personalizados inspirados en admin_material
    titulo_principal = ParagraphStyle(
        'TituloPrincipal',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=12,
        alignment=TA_CENTER,
        textColor=HexColor('#1976d2'),  # Azul Material Design
        fontName='Helvetica-Bold'
    )
    
    titulo_secundario = ParagraphStyle(
        'TituloSecundario',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=8,
        alignment=TA_CENTER,
        textColor=HexColor('#424242'),  # Gris oscuro Material Design
        fontName='Helvetica-Bold'
    )
    
    titulo_boleta = ParagraphStyle(
        'TituloBoleta',
        parent=styles['Heading2'],
        fontSize=13,
        spaceAfter=16,
        alignment=TA_CENTER,
        textColor=HexColor('#d32f2f'),  # Rojo Material Design
        fontName='Helvetica-Bold'
    )
    
    texto_normal = ParagraphStyle(
        'TextoNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        textColor=HexColor('#212121'),  # Negro Material Design
        fontName='Helvetica'
    )
    
    # Construir el contenido del PDF
    story = []
    
    # Encabezado
    story.append(Paragraph("TECNOLÓGICO NACIONAL DE MÉXICO", titulo_principal))
    story.append(Paragraph("INSTITUTO TECNOLÓGICO SUPERIOR DE TLAXCO", titulo_secundario))
    story.append(Paragraph("Depto. de Servicios Escolares", texto_normal))
    story.append(Spacer(1, 20))
    
    # Título de la boleta
    story.append(Paragraph(f"BOLETA DE CALIFICACIONES PERÍODO {periodo.ciclo.upper()}-{periodo.año}", titulo_boleta))
    story.append(Spacer(1, 20))
    
    # Información del alumno
    info_data = [
        ['N° CONTROL', 'NOMBRE', 'CARRERA', 'SEMESTRE', 'PROMEDIO'],
        [
            alumno.matricula,
            f"{alumno.apellido_paterno} {alumno.apellido_materno} {alumno.nombre}",
            alumno.carrera.nombre if alumno.carrera else "",
            str(alumno.semestre),
            f"{alumno.promedio:.2f}"
        ]
    ]
    
    info_table = Table(info_data, colWidths=[3*cm, 6*cm, 5*cm, 2*cm, 2*cm])
    info_table.setStyle(TableStyle([
        # Encabezado
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#e3f2fd')),  # Azul claro Material Design
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#1976d2')),   # Azul Material Design
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        
        # Datos
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), HexColor('#424242')),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 1, HexColor('#e0e0e0')),  # Gris claro Material Design
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
    ]))
    
    story.append(info_table)
    story.append(Spacer(1, 20))
    
    # Tabla de calificaciones
    calificaciones = boleta.get_calificaciones()
    
    if calificaciones.exists():
        # Encabezados de la tabla
        cal_data = [['CLAVE', 'MATERIA', 'NIVEL DE DESEMPEÑO', 'VALORACIÓN NUMÉRICA', 'OPCIÓN', 'CRÉDITOS']]
        
        # Agregar calificaciones
        for calif in calificaciones:
            nivel_desempeno = obtener_nivel_desempeno(calif.calificacion)
            cal_data.append([
                calif.materia.clave,
                calif.materia.nombre,
                nivel_desempeno,
                str(calif.calificacion) if calif.calificacion else "0",
                calif.acreditacion[0] if calif.acreditacion else "O",
                str(calif.creditos)
            ])
        
        cal_table = Table(cal_data, colWidths=[2*cm, 6*cm, 3*cm, 2.5*cm, 1.5*cm, 2*cm])
        cal_table.setStyle(TableStyle([
            # Encabezado
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#f3e5f5')),  # Púrpura claro Material Design
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#7b1fa2')),   # Púrpura Material Design
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            
            # Datos
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), HexColor('#424242')),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, HexColor('#e0e0e0')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            
            # Alineación específica para columnas
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),  # Materia alineada a la izquierda
        ]))
        
        story.append(cal_table)
    else:
        story.append(Paragraph("No hay calificaciones registradas para este período.", texto_normal))
    
    story.append(Spacer(1, 20))
    
    # Leyenda de niveles de desempeño
    leyenda_style = ParagraphStyle(
        'Leyenda',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#616161'),  # Gris medio Material Design
        fontName='Helvetica'
    )
    
    story.append(Paragraph("<b>NIVEL DE DESEMPEÑO:</b> DI=Desempeño insuficiente, S=Suficiente, B=Bueno, N=Notable, E=Excelente", leyenda_style))
    story.append(Spacer(1, 40))
    
    # Firma
    firma_style = ParagraphStyle(
        'Firma',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        textColor=HexColor('#424242'),
        fontName='Helvetica'
    )
    
    story.append(Paragraph("_" * 60, firma_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("<b>LIC. ISRAEL ALLAN MORALES BARRIOS</b>", firma_style))
    story.append(Paragraph("JEFE DE DEPARTAMENTO DE SERVICIOS ESCOLARES", firma_style))
    
    # Construir el PDF
    doc.build(story)
    
    # Preparar la respuesta
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Boleta_{alumno.matricula}_{periodo}.pdf"'
    
    return response


def crear_encabezado(doc, alumno, periodo):
    """Crear el encabezado con logos y título"""
    # Título principal
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = titulo.add_run("TECNOLÓGICO NACIONAL DE MÉXICO\n")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    
    run = titulo.add_run("INSTITUTO TECNOLÓGICO SUPERIOR DE TLAXCO\n")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = True
    
    run = titulo.add_run("Depto. de Servicios Escolares\n\n")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    
    # Título de la boleta
    titulo_boleta = doc.add_paragraph()
    titulo_boleta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = titulo_boleta.add_run(f"BOLETA DE CALIFICACIONES PERÍODO {periodo.ciclo.upper()}-{periodo.año}")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True
    
    doc.add_paragraph()  # Espacio


def crear_info_alumno(doc, alumno, periodo):
    """Crear la tabla con información del alumno"""
    tabla_info = doc.add_table(rows=1, cols=6)
    tabla_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Configurar encabezados
    headers = ["N° CONTROL", "NOMBRE", "CARRERA", "SEMESTRE", "PROMEDIO"]
    row = tabla_info.rows[0]
    
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Arial'
    
    # Agregar datos del alumno
    row_data = tabla_info.add_row()
    datos = [
        alumno.matricula,
        f"{alumno.apellido_paterno} {alumno.apellido_materno} {alumno.nombre}",
        alumno.carrera.nombre if alumno.carrera else "",
        str(alumno.semestre),
        f"{alumno.promedio:.2f}"
    ]
    
    for i, dato in enumerate(datos):
        cell = row_data.cells[i]
        cell.text = dato
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.size = Pt(8)
        run.font.name = 'Arial'
    
    doc.add_paragraph()  # Espacio


def crear_tabla_calificaciones(doc, boleta):
    """Crear la tabla de calificaciones"""
    calificaciones = boleta.get_calificaciones()
    
    if not calificaciones.exists():
        p = doc.add_paragraph("No hay calificaciones registradas para este período.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    
    # Crear tabla
    tabla = doc.add_table(rows=1, cols=6)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Encabezados
    headers = ["CLAVE", "MATERIA", "NIVEL DE DESEMPEÑO", "VALORACIÓN NUMÉRICA", "OPCIÓN", "CRÉDITOS"]
    row = tabla.rows[0]
    
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Arial'
        
        # Color de fondo del encabezado
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
    
    # Agregar calificaciones
    for calif in calificaciones:
        row_data = tabla.add_row()
        
        # Determinar nivel de desempeño
        nivel_desempeno = obtener_nivel_desempeno(calif.calificacion)
        
        datos = [
            calif.materia.clave,
            calif.materia.nombre,
            nivel_desempeno,
            str(calif.calificacion) if calif.calificacion else "0",
            calif.acreditacion[0] if calif.acreditacion else "O",  # Usar primera letra de acreditación
            str(calif.creditos)  # Usar créditos de la calificación, no de la materia
        ]
        
        for i, dato in enumerate(datos):
            cell = row_data.cells[i]
            cell.text = dato
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3, 4, 5] else WD_ALIGN_PARAGRAPH.LEFT
            run = para.runs[0]
            run.font.size = Pt(8)
            run.font.name = 'Arial'
    
    doc.add_paragraph()  # Espacio


def obtener_nivel_desempeno(calificacion):
    """Obtener el nivel de desempeño basado en la calificación"""
    if calificacion is None or calificacion == 0:
        return "DI"
    elif calificacion < 60:
        return "DI"
    elif calificacion < 70:
        return "S"
    elif calificacion < 80:
        return "B"
    elif calificacion < 90:
        return "N"
    else:
        return "E"


def crear_leyenda_desempeno(doc):
    """Crear la leyenda de niveles de desempeño"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = p.add_run("NIVEL DE DESEMPEÑO: ")
    run.font.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Arial'
    
    run = p.add_run("DI=Desempeño insuficiente, S=Suficiente, B=Bueno, N=Notable, E=Excelente")
    run.font.size = Pt(8)
    run.font.name = 'Arial'
    
    doc.add_paragraph()  # Espacio


def crear_pie_firma(doc):
    """Crear el pie con espacio para firma"""
    # Agregar espacios
    for _ in range(3):
        doc.add_paragraph()
    
    # Línea para firma
    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p_firma.add_run("_" * 50)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    
    # Nombre del jefe
    p_nombre = doc.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p_nombre.add_run("LIC. ISRAEL ALLAN MORALES BARRIOS\n")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    
    run = p_nombre.add_run("JEFE DE DEPARTAMENTO DE SERVICIOS ESCOLARES")
    run.font.size = Pt(8)
    run.font.name = 'Arial'


@login_required
def ajax_generar_boleta(request):
    """Vista AJAX para generar boleta"""
    if request.method == 'POST':
        alumno_id = request.POST.get('alumno_id')
        periodo_id = request.POST.get('periodo_id')
        
        if not alumno_id or not periodo_id:
            return JsonResponse({'error': 'Faltan parámetros requeridos'}, status=400)
        
        try:
            alumno = Alumno.objects.get(id=alumno_id)
            periodo = PeriodoEscolar.objects.get(id=periodo_id)
            
            # Crear o obtener la boleta
            boleta, created = Boleta.objects.get_or_create(
                alumno=alumno,
                periodo_escolar=periodo,
                defaults={'generado_por': request.user}
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Boleta generada exitosamente',
                'boleta_id': boleta.id,
                'download_url': f'/procedimientos/boleta/documento/{alumno_id}/{periodo_id}/'
            })
            
        except (Alumno.DoesNotExist, PeriodoEscolar.DoesNotExist):
            return JsonResponse({'error': 'Alumno o período no encontrado'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Método no permitido'}, status=405)