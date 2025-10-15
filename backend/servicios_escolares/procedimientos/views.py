from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib import messages
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.db.models import Q, Count
from django.utils import timezone
from django.http import HttpResponse
from uritemplate import variables
from datos_academicos.models import Alumno, Calificacion, Materia, PeriodoEscolar
from procedimientos.models import Tramite, Bitacora, Proceso
from .models import Tramite, Bitacora
from .forms import TramiteForm
from docxtpl import DocxTemplate
from docx2pdf import convert
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Cm
from docsbuilder.models import Plantilla
from docsbuilder.utils import armar_contexto_para_alumno
from collections import defaultdict
from datetime import datetime
from io import BytesIO
import tempfile
import pythoncom
import os

class TramiteListView(LoginRequiredMixin, ListView):
    model = Tramite
    template_name = 'procedimientos/tramite_list.html'
    context_object_name = 'tramites'
    paginate_by = 10

    def get_queryset(self):
        queryset = Tramite.objects.select_related('alumno', 'plantilla').all().order_by('-fecha_solicitud')

        q = self.request.GET.get('q', '').strip()
        estado = self.request.GET.get('estado', '').strip()
        tipo = self.request.GET.get('tipo', '').strip()

        if q:
            queryset = queryset.filter(
                Q(alumno__nombre__icontains=q) |
                Q(alumno__apellido_paterno__icontains=q) |
                Q(alumno__apellido_materno__icontains=q) |
                Q(alumno__matricula__icontains=q)
            )

        if estado:
            queryset = queryset.filter(estado=estado)

        if tipo:
            queryset = queryset.filter(tipo=tipo)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['tipo_choices'] = Tramite.TIPO_CHOICES
        context['segment'] = 'tramites'
        return context

@login_required
def descargar_boleta(request, matricula):
    """
    Genera y descarga una boleta de calificaciones para un alumno específico en formato PDF
    """
    try:
        # Obtener el alumno
        alumno = get_object_or_404(Alumno, matricula=matricula)
        
        # Obtener el periodo (si se especifica, sino el más reciente)
        periodo_id = request.GET.get('periodo')
        if periodo_id:
            periodo = get_object_or_404(PeriodoEscolar, id=periodo_id)
        else:
            periodo = PeriodoEscolar.objects.filter(activo=True).first()
            if not periodo:
                periodo = PeriodoEscolar.objects.order_by('-fecha_fin').first()
        
        if not periodo:
            messages.error(request, 'No hay períodos escolares disponibles.')
            return redirect('procedimientos:boleta_list')
        
        # Obtener la plantilla "Boleta" (ID 20 según la consulta anterior)
        try:
            plantilla = Plantilla.objects.get(nombre="Boleta")
        except Plantilla.DoesNotExist:
            messages.error(request, 'No se encontró la plantilla "Boleta".')
            return redirect('procedimientos:boleta_list')
        
        # Generar el contexto para la boleta
        from docsbuilder.utils import armar_contexto_para_boleta
        variables = plantilla.variables.all()
        contexto = armar_contexto_para_boleta(alumno, periodo, variables)
        
        # Debug: Verificar el contexto generado
        print("=== DEBUG BOLETA ===")
        print(f"Contexto keys: {list(contexto.keys())}")
        if 'calificaciones' in contexto:
            print(f"Calificaciones: {len(contexto['calificaciones'])} registros")
            if contexto['calificaciones']:
                print(f"Primera: {contexto['calificaciones'][0]}")
        print("=== FIN DEBUG ===")
        
        # Generar el documento usando DocxTemplate
        ruta_plantilla = plantilla.archivo.path
        tpl = DocxTemplate(ruta_plantilla)
        tpl.render(contexto)
        
        # Guardar DOCX a archivo temporal
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tpl.save(tmp_docx.name)
            docx_path = tmp_docx.name

        # Procesar marcadores automáticos como en kardex
        doc = Document(docx_path)
        
        # Buscar marcador <<calificaciones>> para generar tabla automáticamente
        def buscar_parrafo_con_marcador(doc, marcador="calificaciones"):
            import re
            patron = re.compile(r'<<' + re.escape(marcador) + r'>>')
            for p in doc.paragraphs:
                if patron.search(p.text):
                    print(f"Marcador encontrado: {p.text}")  # Debug
                    return p
            print(f"Marcador <<{marcador}>> no encontrado")  # Debug
            return None

        p = buscar_parrafo_con_marcador(doc, "calificaciones")
        if p:
            # Obtener calificaciones del período
            calificaciones = Calificacion.objects.filter(
                alumno=alumno,
                periodo_escolar=periodo
            ).select_related('materia').order_by('materia__clave')
            
            if calificaciones.exists():
                # Crear tabla de calificaciones
                p_element = p._element
                
                # Crear tabla
                tabla = doc.add_table(rows=1, cols=5)
                tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Encabezados
                headers = ["CLAVE", "MATERIA", "CRÉDITOS", "CALIFICACIÓN", "ACREDITACIÓN"]
                row = tabla.rows[0]
                
                for i, header in enumerate(headers):
                    cell = row.cells[i]
                    cell.text = header
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
                    
                    # Color de fondo del encabezado
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), '1B396A')
                    tcPr.append(shd)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                
                # Agregar filas de calificaciones
                for idx, calif in enumerate(calificaciones):
                    row_data = tabla.add_row()
                    
                    # Determinar acreditación
                    acreditacion = "ACREDITADA" if calif.calificacion >= 70 else "NO ACREDITADA"
                    
                    datos = [
                        calif.materia.clave,
                        calif.materia.nombre,
                        str(calif.materia.creditos),
                        str(calif.calificacion),
                        acreditacion
                    ]
                    
                    for i, dato in enumerate(datos):
                        cell = row_data.cells[i]
                        cell.text = dato
                        para = cell.paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(8)
                            run.font.name = 'Arial'
                        
                        # Color alternado para filas
                        color = "FFFFFF" if idx % 2 == 0 else "F2F2F2"
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), color)
                        tcPr.append(shd)
                
                # Insertar tabla en el documento
                p_element.addnext(tabla._tbl)
                
                # Eliminar marcador original
                p._element.getparent().remove(p._element)
        
        # Guardar documento final con tabla generada
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as final_docx:
            doc.save(final_docx.name)
            final_docx_path = final_docx.name
        
        # Usar el documento final para conversión a PDF
        os.remove(docx_path)  # Eliminar el temporal anterior
        docx_path = final_docx_path

        # Convertir a PDF con docx2pdf
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
            pdf_path = tmp_pdf.name

        try:
            pythoncom.CoInitialize()
            convert(docx_path, pdf_path)
            pythoncom.CoUninitialize()
        except Exception as e:
            os.remove(docx_path)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            messages.error(request, f'Error al convertir a PDF: {str(e)}')
            return redirect('procedimientos:boleta_list')

        # Leer PDF generado
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        # Limpiar archivos temporales
        os.remove(docx_path)
        os.remove(pdf_path)

        # Preparar la respuesta para descarga en PDF
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="boleta_{alumno.matricula}_{periodo.ciclo}_{periodo.año}.pdf"'
        
        # Registrar NUEVO trámite cada vez que se descarga
        from procedimientos.models import Tramite, Bitacora
        tramite = Tramite.objects.create(
            alumno=alumno,
            tipo='boleta',
            estado='Procesado',
            fecha_procesado=timezone.now(),
            plantilla=plantilla
        )
        
        # Registrar en bitácora
        Bitacora.objects.create(
            tramite=tramite,
            usuario=request.user,
            accion='Generó boleta PDF',
            comentario=f'Boleta generada para el periodo {periodo.ciclo} {periodo.año}'
        )
        return response
        
    except Exception as e:
        messages.error(request, f'Error al generar la boleta: {str(e)}')
        return redirect('procedimientos:boleta_list')


@login_required
def crear_tramite(request):
    if request.method == 'POST':
        form = TramiteForm(request.POST)
        if form.is_valid():
            tramite = form.save()
            Bitacora.objects.create(
                tramite=tramite,
                usuario=request.user,
                accion="Creó trámite",
                comentario=f"Trámite tipo {tramite.get_tipo_display()} creado."
            )
            return redirect('procedimientos:lista_tramites')
    else:
        form = TramiteForm()
    return render(request, 'procedimientos/crear_tramite.html', {'form': form})

@login_required
def descargar_constancia(request, matricula):
    alumno = get_object_or_404(Alumno, matricula=matricula)
    plantilla = Plantilla.objects.filter(nombre__iexact='constancia').first()

    if not plantilla:
        return HttpResponse("No hay plantilla de constancia configurada.", status=404)

    variables = plantilla.variables.all()
    # Usa tu función para armar el contexto
    contexto = armar_contexto_para_alumno(alumno, variables)
    # Asegura fecha de emisión en contexto si no está
    contexto.setdefault('fecha_emision', datetime.now().strftime('%d/%m/%Y'))

    ruta_plantilla = plantilla.archivo.path
    doc = DocxTemplate(ruta_plantilla)
    doc.render(contexto)

    # Guardar DOCX a archivo temporal
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        doc.save(tmp_docx.name)
        docx_path = tmp_docx.name

    # Convertir a PDF con docx2pdf
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
        pdf_path = tmp_pdf.name

    try:
        pythoncom.CoInitialize()
        convert(docx_path, pdf_path)
        pythoncom.CoUninitialize()
    except Exception as e:
        os.remove(docx_path)
        os.remove(pdf_path)
        return HttpResponse(f"Error al convertir a PDF: {str(e)}", status=500)

    # Leer PDF generado
    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    # Limpiar archivos temporales
    os.remove(docx_path)
    os.remove(pdf_path)

    # Registrar trámite constancia
    Tramite.objects.create(
        alumno=alumno,
        tipo='constancia',
        estado='Procesado',
        fecha_procesado=datetime.now(),
        plantilla=plantilla
    )

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Constancia_{alumno.matricula}.pdf"'
    return response

@login_required
def descargar_kardex(request, matricula):
    alumno = get_object_or_404(Alumno, matricula=matricula)
    plantilla = Plantilla.objects.filter(nombre__iexact='kardex').first()
    if not plantilla:
        return HttpResponse("No hay plantilla de Kardex configurada.", status=404)

    # Obtener materias de la carrera del alumno usando la nueva relación
    from datos_academicos.models import MateriaCarrera
    materias_carrera_relaciones = MateriaCarrera.objects.filter(carrera=alumno.carrera).select_related('materia').order_by('semestre', 'materia__clave')
    
    # Obtener todas las materias universales
    materias_universales = Materia.objects.filter(es_universal=True).order_by('clave')
    
    # Obtener todas las calificaciones del alumno (incluyendo de otras carreras)
    califs = Calificacion.objects.filter(alumno=alumno).order_by('-fecha_registro')
    calif_dict = {c.materia_id: c for c in califs}

    contexto = {
        "nombre": alumno.nombre,
        "apellido_paterno": alumno.apellido_paterno or '',
        "apellido_materno": alumno.apellido_materno or '',
        "matricula": alumno.matricula,
        "semestre": str(alumno.semestre),
        "creditos_aprobados": alumno.creditos_aprobados,
        "creditos_totales": alumno.creditos_totales,
        "carrera": alumno.carrera.nombre,
        "plan_estudio": alumno.plan_estudio.clave if alumno.plan_estudio else '',
        "PROMEDIO": alumno.calcular_promedio(incluir_todas=False),  # Solo materias que cuentan para promedio
        "fecha_emision": datetime.now().strftime('%d/%m/%Y'),
    }

    tpl = DocxTemplate(plantilla.archivo.path)
    tpl.render(contexto)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        tpl.save(tmp_docx.name)
        docx_path = tmp_docx.name

    doc = Document(docx_path)
    # Buscar cualquier marcador con formato <<TEXTO_VARIABLE>>
    def buscar_parrafo_con_marcador(doc):
        import re
        patron = re.compile(r'<<.*?>>')
        for p in doc.paragraphs:
            if patron.search(p.text):
                return p
        return None

    p = buscar_parrafo_con_marcador(doc)
    if not p:
        return HttpResponse("No se encontró el marcador de tabla en la plantilla.", status=400)

    # Separar materias por tipo y semestre usando la nueva relación
    materias_por_semestre = defaultdict(list)
    materias_especialidad = []
    materias_universales_list = []
    materias_actividades = []
    
    # Obtener todas las materias con calificaciones del alumno
    materias_con_calificacion = set(calif_dict.keys())
    
    # Procesar materias de la carrera usando la relación MateriaCarrera
    for relacion in materias_carrera_relaciones:
        materia = relacion.materia
        if materia.tipo == "Especialidad":
            materias_especialidad.append(materia)
        elif materia.tipo == "Universal":
            materias_universales_list.append(materia)
        elif materia.tipo == "Actividad":
            materias_actividades.append(materia)
        else:  # Obligatorias
            # Usar el semestre específico de la relación MateriaCarrera
            semestre = relacion.semestre
            if semestre:
                materias_por_semestre[semestre].append(materia)
    
    # Agregar materias universales
    for materia in materias_universales:
        if materia.id in materias_con_calificacion:
            materias_universales_list.append(materia)
    
    # Agregar materias de otras carreras que el alumno haya cursado
    materias_otras_carreras = []
    for calif_id in materias_con_calificacion:
        materia = Materia.objects.get(id=calif_id)
        # Si la materia no pertenece a la carrera del alumno (no existe relación MateriaCarrera)
        if not MateriaCarrera.objects.filter(materia=materia, carrera=alumno.carrera).exists():
            materias_otras_carreras.append(materia)

    p_element = p._element

    def quitar_bordes_de_tabla(tabla):
        tbl_pr = tabla._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tabla._tbl.insert(0, tbl_pr)
        tbl_borders = OxmlElement('w:tblBorders')
        for borde in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border_el = OxmlElement(f'w:{borde}')
            border_el.set(qn('w:val'), 'nil')  # quitar borde
            tbl_borders.append(border_el)
        tbl_pr.append(tbl_borders)

    def set_cell_background(cell, color_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    # Función para crear tabla con título y filas dadas
    def crear_tabla_kardex(titulo, lista_materias):
        nonlocal p_element

        p_titulo = doc.add_paragraph(titulo)
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_titulo.runs[0]
        run.bold = True
        run.font.family = 'Arial'
        run.font.size = Pt(9)

        tabla = doc.add_table(rows=1, cols=5)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        quitar_bordes_de_tabla(tabla)

        headers = ["Clave", "Materia", "Créditos", "Calificación", "Acreditación"]

        for i, texto in enumerate(headers):
            cell = tabla.rows[0].cells[i]
            cell.text = texto.upper()
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.family = 'Arial'
            run.font.size = Pt(8)

            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1B396A')
            tcPr.append(shd)

        col_widths = [2, 6, 2, 2, 3]
        for i, w in enumerate(col_widths):
            for row in tabla.rows:
                row.cells[i].width = Cm(w)

        for idx, materia in enumerate(lista_materias):
            calif = calif_dict.get(materia.id)
            row_cells = tabla.add_row().cells
            row_cells[0].text = materia.clave
            row_cells[1].text = materia.nombre
            row_cells[2].text = str(materia.creditos)
            row_cells[3].text = str(calif.calificacion) if calif else "—"
            row_cells[4].text = calif.acreditacion if calif else "—"

            color = "FFFFFF" if idx % 2 == 0 else "F2F2F2"
            for cell in row_cells:
                set_cell_background(cell, color)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(8)

        # Insertar título y tabla en el documento
        p_element.addnext(p_titulo._element)
        p_element = p_titulo._element
        p_element.addnext(tabla._tbl)
        p_element = tabla._tbl

    # Crear tablas para semestres normales
    for semestre in sorted(materias_por_semestre.keys()):
        crear_tabla_kardex(f"SEMESTRE {semestre}", materias_por_semestre[semestre])

    # Crear tabla para Especialidad si hay materias
    if materias_especialidad:
        crear_tabla_kardex("MATERIAS DE ESPECIALIDAD", materias_especialidad)
    
    # Crear tabla para materias universales si hay
    if materias_universales_list:
        crear_tabla_kardex("OTRAS", materias_universales_list)
    
    # Crear tabla para materias de otras carreras si hay
    if materias_otras_carreras:
        crear_tabla_kardex("OTRAS ACREDITACIONES", materias_otras_carreras)
    
    # Crear tabla para actividades si hay
    if materias_actividades:
        crear_tabla_kardex("ACTIVIDADES COMPLEMENTARIAS", materias_actividades)

    # Eliminar marcador original
    p._element.getparent().remove(p._element)

    # Guardar documento final
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as final_docx:
        doc.save(final_docx.name)
        final_path = final_docx.name

    # Convertir a PDF
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
        pdf_path = tmp_pdf.name

    try:
        import pythoncom
        from docx2pdf import convert
        pythoncom.CoInitialize()
        convert(final_path, pdf_path)
        pythoncom.CoUninitialize()
    except Exception as e:
        return HttpResponse(f"Error al convertir a PDF: {str(e)}", status=500)

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    os.remove(docx_path)
    os.remove(final_path)
    os.remove(pdf_path)

    Tramite.objects.create(
        alumno=alumno,
        tipo='kardex',
        estado='Procesado',
        fecha_procesado=datetime.now(),
        plantilla=plantilla
    )

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Kardex_{alumno.matricula}.pdf"'
    return response

# Auxiliar
def actualizar_creditos_alumno(alumno):
    # Obtener materias a través de la relación MateriaCarrera
    materias_carrera = MateriaCarrera.objects.filter(carrera=alumno.carrera).select_related('materia')
    materias = [mc.materia for mc in materias_carrera]
    
    califs = Calificacion.objects.filter(alumno=alumno)
    calif_dict = {c.materia_id: c for c in califs}

    total = sum(m.creditos for m in materias)
    aprobados = sum(
        m.creditos for m in materias
        if (c := calif_dict.get(m.id)) and c.calificacion and float(c.calificacion) >= 6.0
    )

    alumno.creditos_totales = total
    alumno.creditos_aprobados = aprobados
    alumno.save(update_fields=['creditos_totales', 'creditos_aprobados'])

class ProcesoListView(LoginRequiredMixin, ListView):
    model = Proceso
    template_name = 'procedimientos/proceso_list.html'
    context_object_name = 'procesos'
    paginate_by = 10

    def get_queryset(self):
        return Proceso.objects.filter(activo=True).order_by('-fecha_inicio')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['segment'] = 'procesos'
        return context

@login_required
def dashboard_tramites(request):
    """Vista del dashboard de trámites con KPIs, gráficos y acciones principales"""
    
    # KPIs de trámites
    total_tramites = Tramite.objects.count()
    tramites_pendientes = Tramite.objects.filter(estado='En proceso').count()
    tramites_completados = Tramite.objects.filter(estado='Completado').count()
    tramites_rechazados = Tramite.objects.filter(estado='Rechazado').count()
    
    # Calcular porcentajes
    porcentaje_completados = (tramites_completados / total_tramites * 100) if total_tramites > 0 else 0
    porcentaje_pendientes = (tramites_pendientes / total_tramites * 100) if total_tramites > 0 else 0
    
    # Datos para gráfico de trámites por tipo
    datos_tipos = (
        Tramite.objects
        .values('tipo')
        .annotate(cantidad=Count('id'))
        .order_by('-cantidad')
    )
    tipos_tramites = [item['tipo'] for item in datos_tipos]
    cantidades_tipos = [item['cantidad'] for item in datos_tipos]
    
    # Datos para gráfico de trámites por estado
    datos_estados = (
        Tramite.objects
        .values('estado')
        .annotate(cantidad=Count('id'))
        .order_by('-cantidad')
    )
    estados_tramites = [item['estado'] for item in datos_estados]
    cantidades_estados = [item['cantidad'] for item in datos_estados]
    
    # Trámites recientes (últimos 10)
    tramites_recientes = Tramite.objects.order_by('-fecha_solicitud')[:10]
    
    # Procesos activos
    procesos_activos = Proceso.objects.filter(activo=True).count()
    
    context = {
        'segment': 'dashboard_tramites',
        'total_tramites': total_tramites,
        'tramites_pendientes': tramites_pendientes,
        'tramites_completados': tramites_completados,
        'tramites_rechazados': tramites_rechazados,
        'porcentaje_completados': round(porcentaje_completados, 1),
        'porcentaje_pendientes': round(porcentaje_pendientes, 1),
        'tipos_tramites': tipos_tramites,
        'cantidades_tipos': cantidades_tipos,
        'estados_tramites': estados_tramites,
        'cantidades_estados': cantidades_estados,
        'tramites_recientes': tramites_recientes,
        'procesos_activos': procesos_activos,
    }
    
    return render(request, 'procedimientos/dashboard_tramites.html', context)

